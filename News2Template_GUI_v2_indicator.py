# -*- coding: utf-8 -*-
import os
import re
import csv
import threading
import queue
import webbrowser
from typing import Optional, Dict, List, Tuple

import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# Optional drag & drop
_DND_AVAILABLE = False
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    _DND_AVAILABLE = True
except Exception:
    _DND_AVAILABLE = False

# Doc processing deps
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

_HAVE_PDFMINER = False
_HAVE_TEXTRACT = False
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    _HAVE_PDFMINER = True
except Exception:
    _HAVE_PDFMINER = False

try:
    import textract
    _HAVE_TEXTRACT = True
except Exception:
    _HAVE_TEXTRACT = False

CONFIG = {
    "header_text": "广东第二师范学院音乐学院 | 学生工作年度考核材料汇总",
    # default if no match
    "indicator_fallback": "（待匹配）",
    "default_org": "音乐学院学生工作办公室",
    "allowed_exts": [".txt", ".docx", ".doc", ".pdf"],
    "output_dirname": "output_docs",
}

# ---------------- 指标体系（关键词规则） ----------------
# 每项：("代码 名称", [关键词...])
INDICATOR_RULES: List[Tuple[str, List[str]]] = [
    ("1.1 领导重视", ["年度工作计划", "工作计划", "工作总结", "学生工作会议", "党组织会议", "专题座谈会", "座谈会", "走访宿舍", "宿舍走访", "倾听学生", "党政领导", "专题会议"]),
    ("1.2 学生事务管理", ["学风建设", "学业帮扶", "突发事件", "应急预案", "值班", "安全教育", "防诈骗", "消防", "交通安全", "劳动教育", "宿舍检查", "纪律教育", "安全检查", "稳定工作"]),
    ("1.3 资助育人", ["资助", "经济困难", "奖助", "认定", "受助学生", "资助政策", "诚信教育", "感恩教育", "励志", "成长成才"]),
    ("1.4 国防教育", ["征兵", "国防教育", "军训", "入伍", "退役", "服兵役", "应征"]),
    ("1.5 心理健康教育", ["心理健康", "心理访谈", "心理排查", "研判", "心理测评", "心理工作站", "重点关注学生", "心理教育"]),
    ("1.6 体育育人", ["阳光体育", "体育运动", "长跑", "体测", "运动会", "体育课程"]),
    ("2 就业创业工作", ["就业", "签约率", "招聘会", "双选会", "就业指导", "创业", "实习基地", "访企拓岗"]),
    ("3 共青团工作", ["共青团", "团委", "团支部", "团员", "主题团日", "青马班", "团课", "团学骨干", "团学组织"]),
    ("4.1 辅导员队伍项目与获奖", ["辅导员", "副书记", "论文", "项目立项", "大赛获奖", "学生工作类论文", "成果获奖"]),
    ("4.2 学生队伍获奖", ["学生获奖", "比赛获奖", "大赛", "荣誉称号", "省赛", "国赛", "校赛"]),
    ("5.1 扣分项（事故/违纪）", ["政治安全事件", "安全事故", "立案处理", "通报批评", "记过处分", "违纪", "处分"]),
]

def match_indicator(text: str) -> Tuple[str, List[str], int]:
    """
    返回 (最佳指标, 命中关键词列表, 得分)
    简单策略：按关键字命中计分，取最高；平分时按列表先后优先。
    """
    txt = text or ""
    best = (CONFIG["indicator_fallback"], [], 0)
    for code_name, kws in INDICATOR_RULES:
        hits = []
        score = 0
        for kw in kws:
            if kw in txt:
                hits.append(kw)
                score += 1
        # 兼容近义表达（正则微扩展）
        if re.search(r"安全(教育|检查)", txt):
            if "安全教育" not in hits and "安全检查" not in hits:
                hits.append("安全*")
                score += 1
        if score > best[2]:
            best = (code_name, hits, score)
    return best

# ---------------- 文本读取与结构化 ----------------
def read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read().strip()

def read_docx(path: str) -> str:
    doc = DocxDocument(path)
    paras = [p.text for p in doc.paragraphs]
    return "\n".join([p for p in paras if p is not None]).strip()

def read_doc(path: str) -> Optional[str]:
    if not _HAVE_TEXTRACT:
        return None
    try:
        raw = textract.process(path)
        if isinstance(raw, bytes):
            return raw.decode("utf-8", errors="ignore").strip()
        return str(raw)
    except Exception:
        return None

def read_pdf(path: str) -> Optional[str]:
    if not _HAVE_PDFMINER:
        return None
    try:
        txt = pdf_extract_text(path)
        return (txt or "").strip()
    except Exception:
        return None

def extract_first_line_as_title(text: str) -> str:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if not lines:
        return "（未命名活动）"
    return lines[0][:80]

def extract_date(text: str) -> Optional[str]:
    m = re.search(r'(\d{4}年\d{1,2}月\d{1,2}日)', text)
    return m.group(1) if m else None

def extract_location(text: str) -> Optional[str]:
    m = re.search(r'在([^\n，。；；、]{2,30}?)(举行|举办|召开|隆重举行|开展|进行)', text)
    if m:
        return m.group(1).strip()
    m = re.search(r'(地点[:：为]\s*([^\n，。；]{2,30}))', text)
    if m:
        return m.group(2).strip()
    m = re.search(r'([^\s，。；]{2,20}(中心|学院|礼堂|体育馆|广场|音乐厅|会堂|展馆|研修院|文化馆))', text)
    if m:
        return m.group(1).strip()
    return None

def extract_org(text: str) -> Optional[str]:
    m = re.search(r'(主办|承办|协办)[单机]位[：:]\s*([^\n，。；]{2,40})', text)
    if m:
        return m.group(2).strip()
    m = re.search(r'([^\n，。；]{2,30}(学院|系|处|委|中心))[^\n]{0,10}(组织|参与|承办|协办|以精彩|带队|作为)', text)
    if m:
        return m.group(1).strip()
    return None

def extract_participant_count(text: str) -> Optional[str]:
    m = re.search(r'(参与|参加|到场|观众|师生|学生)[^\n，。；]{0,10}?(\d{1,4})(人|名)', text)
    if m:
        return m.group(2) + "人"
    m = re.search(r'(\d{1,4})(人|名)', text)
    if m:
        return m.group(1) + "人"
    return None

def summarize_overview(text: str) -> str:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    body = "\n".join(lines[1:]) if len(lines) > 1 else text
    sentences = re.split(r'[。！？!?\n]', body)
    sentences = [s.strip() for s in sentences if s.strip()]
    overview = "。".join(sentences[:3])[:200]
    return (overview + "。") if overview else "（活动概述待补充）"

def summarize_effect(text: str) -> str:
    pieces: List[str] = []
    if re.search(r'感谢信|表扬|致谢', text):
        pieces.append("收到相关部门感谢信或表扬。")
    if re.search(r'好评|广泛关注|反响热烈|积极评价', text):
        pieces.append("活动获得积极评价，反响良好。")
    if re.search(r'媒体|公众号|报道|推送|阅读量', text):
        pieces.append("形成宣传报道与传播影响。")
    if not pieces:
        pieces.append("提升学生社会责任感与服务意识，形成良好育人效果。")
    return " ".join(pieces)

# ---------------- 生成文档 ----------------
def create_doc(activity: Dict[str, str], indicator_code: str, out_path: str):
    doc = DocxDocument()
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = CONFIG["header_text"]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.style.font.name = '宋体'
    header.style.font.size = Pt(10)

    footer = section.footer.paragraphs[0]
    footer.text = "第 {PAGE} 页，共 {NUMPAGES} 页"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.style.font.name = '宋体'
    footer.style.font.size = Pt(10)

    doc.add_heading(activity.get("活动名称", "（未命名活动）"), level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows = [
        ("活动名称", activity.get("活动名称", "")),
        ("活动时间", activity.get("活动时间", "")),
        ("活动地点", activity.get("活动地点", "")),
        ("主办/承办单位", activity.get("主办/承办单位", "")),
        ("参与人数", activity.get("参与人数", "")),
        ("对应考核指标", indicator_code or CONFIG["indicator_fallback"]),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v

    doc.add_paragraph("\n【活动概述】").bold = True
    doc.add_paragraph(activity.get("活动概述", ""))

    doc.add_paragraph("\n【成效与佐证】").bold = True
    doc.add_paragraph(activity.get("成效与佐证", ""))

    doc.add_paragraph("\n【图片佐证】（点击下方方框插入图片）").bold = True
    doc.add_paragraph("□ 活动照片1（说明）\n□ 活动照片2（说明）\n□ 海报/媒体报道/感谢信等佐证材料")

    doc.save(out_path)

def parse_article(text: str) -> Dict[str, str]:
    title = extract_first_line_as_title(text)
    date = extract_date(text) or "（日期待补充）"
    loc = extract_location(text) or "（地点待补充）"
    org = extract_org(text) or CONFIG["default_org"]
    ppl = extract_participant_count(text) or "（人数待补充）"
    overview = summarize_overview(text)
    effect = summarize_effect(text)
    return {
        "活动名称": title,
        "活动时间": date,
        "活动地点": loc,
        "主办/承办单位": org,
        "参与人数": ppl,
        "活动概述": overview,
        "成效与佐证": effect,
    }

def load_text_from_file(path: str) -> Optional[str]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        return read_txt(path)
    if ext == ".docx":
        return read_docx(path)
    if ext == ".doc":
        return read_doc(path)
    if ext == ".pdf":
        return read_pdf(path)
    return None

# ---------------- 处理流程 ----------------
def process_folder(input_dir: str, log: tk.Text) -> str:
    output_dir = os.path.join(input_dir, CONFIG["output_dirname"])
    os.makedirs(output_dir, exist_ok=True)

    inputs = []
    for name in os.listdir(input_dir):
        base, ext = os.path.splitext(name)
        if ext.lower() in CONFIG["allowed_exts"]:
            inputs.append(os.path.join(input_dir, name))

    report_rows = [["文件名", "活动名称", "时间", "地点", "单位", "人数", "匹配指标", "匹配关键词", "输出文件", "备注"]]
    count_ok = 0

    for ipath in sorted(inputs):
        log_safe(log, f"处理：{os.path.basename(ipath)}")
        text = load_text_from_file(ipath)
        if not text or len(text.strip()) == 0:
            note = []
            ext = os.path.splitext(ipath)[1].lower()
            if ext == ".pdf" and not _HAVE_PDFMINER:
                note.append("缺少依赖：pdfminer.six")
            if ext == ".doc" and not _HAVE_TEXTRACT:
                note.append("缺少依赖：textract/antiword")
            report_rows.append([os.path.basename(ipath), "", "", "", "", "", "", "", "", "无法读取文本；" + "；".join(note)])
            log_safe(log, f"  × 无法读取文本（{';'.join(note) or '请检查文件是否为扫描件或损坏'}）")
            continue

        try:
            act = parse_article(text)
            code, hits, score = match_indicator(text)
            out_name = f"{act['活动名称'][:30]}_模板版.docx"
            out_name = "".join(ch for ch in out_name if ch not in '\\/:*?\"<>|')
            out_path = os.path.join(output_dir, out_name)
            create_doc(act, code, out_path)

            report_rows.append([
                os.path.basename(ipath),
                act["活动名称"],
                act["活动时间"],
                act["活动地点"],
                act["主办/承办单位"],
                act["参与人数"],
                code,
                "、".join(hits),
                out_path,
                "OK"
            ])
            log_safe(log, f"  ✓ 匹配指标：{code or '（待匹配）'} | 生成：{out_name}")
            count_ok += 1
        except Exception as e:
            report_rows.append([os.path.basename(ipath), "", "", "", "", "", "", "", "", f"ERROR: {e}"])
            log_safe(log, f"  × 处理出错：{e}")

    report_path = os.path.join(output_dir, "report.csv")
    with open(report_path, "w", encoding="utf-8-sig", newline="") as f:
        csv.writer(f).writerows(report_rows)
    log_safe(log, f"完成：成功 {count_ok} / 共 {len(inputs)}。结果清单：{report_path}")
    return output_dir

# ---------------- GUI ----------------
def log_safe(text_widget: tk.Text, msg: str):
    text_widget.configure(state="normal")
    text_widget.insert("end", msg + "\n")
    text_widget.see("end")
    text_widget.configure(state="disabled")
    text_widget.update_idletasks()

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("新闻稿 → 绩效考核模板（GUI v2 指标自动匹配）")

        if _DND_AVAILABLE and isinstance(self.root, TkinterDnD.Tk):
            self.root.drop_target_register(DND_FILES)
            self.root.dnd_bind('<<Drop>>', self.on_drop)

        top = ttk.Frame(root, padding=10)
        top.pack(fill="x")

        self.path_var = tk.StringVar()
        ttk.Label(top, text="输入文件夹：").pack(side="left")
        self.entry = ttk.Entry(top, textvariable=self.path_var, width=60)
        self.entry.pack(side="left", padx=5)
        ttk.Button(top, text="选择文件夹", command=self.select_folder).pack(side="left", padx=5)
        ttk.Button(top, text="开始处理", command=self.start_process).pack(side="left", padx=5)

        info = ttk.Frame(root, padding=(10, 0))
        info.pack(fill="x")
        dnd_msg = "支持拖拽文件夹到窗口" if _DND_AVAILABLE else "（如需拖拽，请安装 tkinterdnd2；当前已提供“选择文件夹”按钮）"
        ttk.Label(info, text=f"支持格式：.txt .docx .doc .pdf | {dnd_msg}").pack(anchor="w")
        ttk.Label(info, text="指标自动匹配：基于关键词规则，支持1.1~5.1条目；结果写入report.csv。").pack(anchor="w")

        self.log = tk.Text(root, height=18, state="disabled")
        self.log.pack(fill="both", expand=True, padx=10, pady=10)

        bottom = ttk.Frame(root, padding=10)
        bottom.pack(fill="x")
        ttk.Button(bottom, text="打开输出文件夹", command=self.open_output).pack(side="left")
        ttk.Button(bottom, text="使用说明", command=self.open_help).pack(side="left", padx=5)

        self.output_dir = None

    def select_folder(self):
        path = filedialog.askdirectory()
        if path:
            self.path_var.set(path)

    def on_drop(self, event):
        paths = self.root.splitlist(event.data)
        if not paths:
            return
        self.path_var.set(paths[0])

    def start_process(self):
        input_dir = self.path_var.get().strip()
        if not input_dir or not os.path.isdir(input_dir):
            messagebox.showwarning("提示", "请选择有效的文件夹")
            return

        self.log.configure(state="normal")
        self.log.delete("1.0", "end")
        self.log.configure(state="disabled")
        log_safe(self.log, f"开始处理：{input_dir}")

        t = threading.Thread(target=self._worker, args=(input_dir,), daemon=True)
        t.start()

    def _worker(self, input_dir: str):
        try:
            self.output_dir = process_folder(input_dir, self.log)
            log_safe(self.log, "处理完成。")
        except Exception as e:
            log_safe(self.log, f"处理失败：{e}")

    def open_output(self):
        if not self.output_dir or not os.path.isdir(self.output_dir):
            messagebox.showinfo("提示", "还没有生成输出文件夹")
            return
        try:
            if os.name == "nt":
                os.startfile(self.output_dir)
            elif os.name == "posix":
                import sys, subprocess
                subprocess.call(["open" if sys.platform == "darwin" else "xdg-open", self.output_dir])
            else:
                webbrowser.open(self.output_dir)
        except Exception:
            webbrowser.open(self.output_dir)

    def open_help(self):
        messagebox.showinfo("使用说明", 
            "1) 把 .txt/.docx/.doc/.pdf 新闻稿放到一个文件夹；\n"
            "2) 选择该文件夹或将其拖入窗口；\n"
            "3) 点击“开始处理”；\n"
            "4) 结果在该文件夹下的 output_docs/ 中，包括 report.csv；\n"
            "5) 工具基于关键词规则自动匹配 1.1~5.1 指标，可在生成的 Word 内人工微调。"
        )

def main():
    if _DND_AVAILABLE:
        app_root = TkinterDnD.Tk()
    else:
        app_root = tk.Tk()
    App(app_root)
    app_root.geometry("800x540")
    app_root.mainloop()

if __name__ == "__main__":
    main()
